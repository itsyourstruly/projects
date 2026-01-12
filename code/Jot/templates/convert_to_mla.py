#!/usr/bin/env python3
"""
Direct conversion from text to MLA-formatted DOCX
Bypasses Pandoc to ensure proper paragraph structure
"""

import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def convert_text_to_mla_docx(txt_path, docx_path):
    """Convert a text file directly to MLA-formatted DOCX"""
    try:
        # Read the text file
        with open(txt_path, 'r') as f:
            lines = f.readlines()

        # Create new document
        doc = Document()

        # Set up margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Extract last name from first line for header
        last_name = "LastName"  # Default
        first_non_empty_line = None
        for line in lines:
            if line.strip():
                first_non_empty_line = line.strip()
                break

        if first_non_empty_line:
            # Try to extract last name (last word in the name line)
            name_parts = first_non_empty_line.split()
            if name_parts:
                last_name = name_parts[-1]  # Take the last word as last name

        # Add header with last name and page number
        section = sections[0]
        header = section.header

        # Get or create header paragraph
        if len(header.paragraphs) > 0:
            header_para = header.paragraphs[0]
        else:
            header_para = header.add_paragraph()

        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Clear any existing content
        header_para.clear()

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Add the last name
        name_run = header_para.add_run(f"{last_name} ")
        name_run.font.name = 'Times New Roman'
        name_run.font.size = Pt(12)

        # Add page number using a simpler field structure
        # This creates a proper PAGE field that Word can update
        page_run = header_para.add_run()
        page_run.font.name = 'Times New Roman'
        page_run.font.size = Pt(12)

        # Create field character elements
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '

        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')

        # Add text element showing "1" as default page number
        t_elem = OxmlElement('w:t')
        t_elem.text = '1'

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        # Append all elements to the run
        page_run._r.append(fldChar_begin)
        page_run._r.append(instrText)
        page_run._r.append(fldChar_separate)
        page_run._r.append(t_elem)
        page_run._r.append(fldChar_end)

        print(f"Added header: {last_name} [PAGE]")

        # Process lines and add as paragraphs
        para_index = 0
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            # Skip empty lines but count them
            if not line:
                i += 1
                continue

            # Add paragraph with content
            para = doc.add_paragraph(line)

            # Apply formatting based on position
            # First 4 non-empty lines: header
            if para_index < 4:
                apply_format(para, 'header')
                print(f"Line {i} (para {para_index}): Header - {line[:40]}")
            # 5th non-empty line: title
            elif para_index == 4:
                apply_format(para, 'title')
                print(f"Line {i} (para {para_index}): TITLE (centered) - {line[:40]}")
            # "Works Cited"
            elif line.strip() == "Works Cited":
                apply_format(para, 'works_cited')
                print(f"Line {i} (para {para_index}): Works Cited")
            # Citations (contains italics or specific keywords)
            elif '*' in line or any(kw in line for kw in ['vol.', 'pp.', '."']):
                apply_format(para, 'citation')
                print(f"Line {i} (para {para_index}): Citation")
            # Body paragraphs
            else:
                apply_format(para, 'body')
                print(f"Line {i} (para {para_index}): Body")

            para_index += 1
            i += 1

        # Save the document
        doc.save(docx_path)
        print(f"Successfully created {docx_path}")
        return True

    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return False

def apply_format(para, style):
    """Apply MLA formatting to a paragraph"""
    # Times New Roman 12pt
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # Double spacing
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    if style == 'header':
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.first_line_indent = Inches(0)
        para.paragraph_format.left_indent = Inches(0)
    elif style == 'title':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Inches(0)
        para.paragraph_format.left_indent = Inches(0)
    elif style == 'works_cited':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Inches(0)
        para.paragraph_format.left_indent = Inches(0)
        if para.runs:
            para.runs[0].bold = True
    elif style == 'citation':
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
    else:  # body
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.first_line_indent = Inches(0.5)
        para.paragraph_format.left_indent = Inches(0)

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python3 convert_to_mla.py <input.txt> <output.docx>")
        sys.exit(1)

    success = convert_text_to_mla_docx(sys.argv[1], sys.argv[2])
    sys.exit(0 if success else 1)

